from openai import OpenAI
import os
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from tkinter.ttk import Progressbar
import time
from dotenv import load_dotenv
from docx.oxml import parse_xml

# 加载环境变量
load_dotenv()

class DocumentProcessor:
    def __init__(self, translator):
        self.translator = translator
        self.processed_elements = 0
        self.total_elements = 0
        self.text_buffer = []  # 添加文本缓冲区
        self.buffer_limit = 1000  # 设置缓冲区字符限制
        self.min_chars = 300  # 增加到300个字符

    def count_translatable_elements(self, doc):
        """计算文档中可翻译元素的总数"""
        count = 0
        # 计算段落数
        for para in doc.paragraphs:
            if para.text.strip():  # 只计算非空段落
                count += 1
                
        # 计算表格中的单元格数
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():  # 只计算非空单元格
                        count += 1
                        
        # 计算文本框数
        for shape in doc.inline_shapes:
            if hasattr(shape, 'text_frame') and shape.text_frame.text.strip():
                count += 1
                
        # 计算页眉页脚数
        for section in doc.sections:
            # 计算页眉段落
            if section.header:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        count += 1
            # 计算页脚段落
            if section.footer:
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        count += 1
                        
        return count

    def flush_buffer(self, target_language):
        """翻译并清空缓冲区"""
        if not self.text_buffer:
            return {}
            
        # 将缓冲区文本合并，用特殊标记分隔
        combined_text = "\n---SPLIT---\n".join(text for text, _ in self.text_buffer)
        translated_text = self.translator.translate_text(combined_text, target_language)
        
        if translated_text:
            # 分割翻译结果
            translations = translated_text.split("\n---SPLIT---\n")
            # 创建原文到译文的映射
            translation_map = {}
            for (original, _), translation in zip(self.text_buffer, translations):
                translation_map[original] = translation.strip()
        else:
            translation_map = {text: text for text, _ in self.text_buffer}
            
        self.text_buffer.clear()
        return translation_map

    def add_to_buffer(self, text, context):
        """添加文本到缓冲区"""
        if text.strip():
            self.text_buffer.append((text, context))
        
        # 如果缓冲区文本总长度超过限制，则执行翻译
        if sum(len(text) for text, _ in self.text_buffer) >= self.buffer_limit:
            return True
        return False

    def translate_paragraphs(self, paragraphs, new_doc, target_language, preserve_format=True, start_index=0):
        """批量翻译段落"""
        translation_map = {}
        
        for i, para in enumerate(paragraphs[start_index:], start=start_index):
            text = para.text.strip()
            if not text:
                new_doc.add_paragraph()
                continue
                
            # 如果文本长度小于阈值且不是最后一个段落，尝试合并
            if len(text) < self.min_chars and i < len(paragraphs) - 1:
                should_translate = self.add_to_buffer(text, para)
            else:
                # 如果文本较长，直接添加并触发翻译
                should_translate = self.add_to_buffer(text, para)
                
            if should_translate:
                translation_map.update(self.flush_buffer(target_language))
                
            # 如果当前段落的翻译已经完成
            if text in translation_map:
                new_para = new_doc.add_paragraph()
                if preserve_format:
                    try:
                        new_para.style = para.style
                    except:
                        pass
                new_para.text = translation_map[text]
                self.processed_elements += 1
                
        # 处理缓冲区中剩余的文本
        if self.text_buffer:
            translation_map.update(self.flush_buffer(target_language))
            for text, para in self.text_buffer:
                new_para = new_doc.add_paragraph()
                if preserve_format:
                    try:
                        new_para.style = para.style
                    except:
                        pass
                new_para.text = translation_map.get(text, text)
                self.processed_elements += 1

    def translate_table(self, source_table, new_doc, target_language, preserve_format=True):
        """翻译表格内容"""
        try:
            rows = len(source_table.rows)
            cols = len(source_table.columns) if source_table.columns else len(source_table.rows[0].cells)
            
            # 在当前位置添加表格
            new_table = new_doc.add_table(rows=rows, cols=cols)
            new_table._element.append(parse_xml(r'<w:tblGrid xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            
            if preserve_format:
                try:
                    new_table.style = source_table.style
                except:
                    pass

            # 使用位置标记收集表格文本
            cell_contents = []
            for i, row in enumerate(source_table.rows):
                for j, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if text:
                        cell_contents.append({
                            'text': text,
                            'row': i,
                            'col': j,
                            'marker': f"[CELL_{i}_{j}]"
                        })

            # 如果没有需要翻译的内容，直接返回
            if not cell_contents:
                return

            # 构建带标记的文本
            marked_text = "\n".join(f"{item['marker']}{item['text']}" for item in cell_contents)
            
            # 翻译带标记的文本
            translated_text = self.translator.translate_text(marked_text, target_language)
            
            if translated_text:
                # 解析翻译结果
                translations = {}
                current_text = ""
                current_marker = None
                
                # 分行处理翻译结果
                for line in translated_text.split('\n'):
                    line = line.strip()
                    if not line:
                        continue
                        
                    # 检查是否找到标记
                    marker_found = False
                    for cell in cell_contents:
                        if line.startswith(cell['marker']):
                            # 如果有未处理的文本，保存它
                            if current_marker:
                                translations[current_marker] = current_text.strip()
                            
                            # 开始新的文本
                            current_marker = cell['marker']
                            current_text = line[len(current_marker):]
                            marker_found = True
                            break
                    
                    if not marker_found and current_marker:
                        # 继续累积当前文本
                        current_text += "\n" + line
                
                # 保存最后一个翻译
                if current_marker:
                    translations[current_marker] = current_text.strip()
                
                # 填充翻译后的表格
                for cell in cell_contents:
                    marker = cell['marker']
                    if marker in translations:
                        new_table.cell(cell['row'], cell['col']).text = translations[marker]
                    else:
                        # 如果没有找到翻译，使用原文
                        new_table.cell(cell['row'], cell['col']).text = cell['text']
                    self.processed_elements += 1

                # 添加一个空段落来分隔表格
                new_doc.add_paragraph()
                    
        except Exception as e:
            print(f"处理表格时出错: {str(e)}")
            # 如果表格处理失败，添加错误信息
            new_doc.add_paragraph("【表格处理失败，原始内容如下】")
            for cell in cell_contents:
                new_doc.add_paragraph(f"行{cell['row']+1}列{cell['col']+1}: {cell['text']}")

    def translate_text_frame(self, source_shape, new_doc, target_language):
        """翻译文本框内容"""
        try:
            if hasattr(source_shape, 'text_frame'):
                text = source_shape.text_frame.text.strip()
                if text:
                    # 翻译文本
                    translated_text = self.translator.translate_text(
                        text,
                        target_language
                    )
                    if translated_text:
                        # 添加一个分隔线表示这是文本框内容
                        new_doc.add_paragraph('─' * 50)
                        # 添加文本框标识
                        new_doc.add_paragraph('【文本框容】').bold = True
                        # 添加翻译后的文本
                        new_para = new_doc.add_paragraph(translated_text)
                        # 添加结束分隔线
                        new_doc.add_paragraph('─' * 50)
                        new_doc.add_paragraph()  # 添加空行
                        self.processed_elements += 1
                        return True
        except Exception as e:
            print(f"处理文本框时出错: {str(e)}")
        return False

    def translate_section(self, section, new_section, target_language):
        """翻译页眉页脚"""
        try:
            # 翻译页眉
            if section.header.paragraphs:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        translated_text = self.translator.translate_text(
                            para.text,
                            target_language
                        )
                        if translated_text:
                            # 确保新文档的页眉有足够的段落
                            while len(new_section.header.paragraphs) <= 0:
                                new_section.header.add_paragraph()
                            new_section.header.paragraphs[0].text = translated_text
                            self.processed_elements += 1

            # 翻译页脚
            if section.footer.paragraphs:
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        translated_text = self.translator.translate_text(
                            para.text,
                            target_language
                        )
                        if translated_text:
                            # 确保新文档的页脚有足够的段落
                            while len(new_section.footer.paragraphs) <= 0:
                                new_section.footer.add_paragraph()
                            new_section.footer.paragraphs[0].text = translated_text
                            self.processed_elements += 1

        except Exception as e:
            print(f"处理页眉页脚时出错: {str(e)}")

class DocTranslator:
    def __init__(self):
        # 从环境变量获取所有 API keys
        self.api_keys = []
        self.current_key_index = 0
        
        # 获取所有配置的API keys
        i = 1
        while True:
            key = os.getenv(f'X_AI_API_KEY_{i}')
            if key:
                self.api_keys.append(key)
                i += 1
            else:
                break
                
        if not self.api_keys:
            raise ValueError("未找到 API key，请在 .env 文件中设置 X_AI_API_KEY_1, X_AI_API_KEY_2 等")
            
        # 为每个 API key 创建一个客户端
        self.clients = []
        for key in self.api_keys:
            client = OpenAI(
                api_key=key,
                base_url="https://api.deepseek.com"
            )
            self.clients.append(client)
        
        self.current_key_index = 0  # 当前使用的API key索引
        
        # 支持的语言字典
        self.supported_languages = {
            "简体中文": "Chinese",
            "繁體中文": "Traditional Chinese",
            "日本語": "Japanese",
            "English": "English",
            "Español": "Spanish",
            "Français": "French",
            "Deutsch": "German",
            "한국어": "Korean",
            "Русский": "Russian",
            "Italiano": "Italian"
        }
    
    def get_next_client(self):
        """获取下一个API客户端"""
        self.current_key_index = (self.current_key_index + 1) % len(self.clients)
        return self.current_key_index
        
    def translate_text(self, text, target_language):
        try:
            # 使用当前客户端
            client_index = self.current_key_index
            
            # 使用选定的客户端发送请求
            if target_language == "English":
                prompt = f"Please translate the following text to English, maintaining professionalism and accuracy:\n\n{text}"
            elif target_language == "Japanese":
                prompt = f"以下のテキストを日本語に翻訳してください。専門性と正確性を保ちながら翻訳してください：\n\n{text}"
            else:
                prompt = f"请将以下文本翻译成{target_language}，保持专业性和准确性：\n\n{text}"

            completion = self.clients[client_index].chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {
                        "role": "system",
                        "content": f"You are a professional translator. Translate the text to {target_language} without adding any additional information or explanations."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=1.3
            )
            return completion.choices[0].message.content
            
        except Exception as e:
            print(f"翻译出错: {str(e)}")
            if "429" in str(e):
                print(f"API Key {client_index + 1} 触发速率限制，切换到下一个 key...")
                # 切换到下一个 API key
                self.get_next_client()
                # 递归重试，使用新的 key
                return self.translate_text(text, target_language)
            return None

class TranslatorGUI:
    def __init__(self):
        try:
            self.translator = DocTranslator()
        except ValueError as e:
            messagebox.showerror("错误", str(e))
            raise
            
        self.window = tk.Tk()
        self.window.title("多语言文档翻译器")
        self.window.geometry("600x400")
        
        self.is_paused = False
        self.translation_start_time = None
        self.processed_paragraphs = 0
        
        # 创建界面元素
        self.setup_gui()
        
    def setup_gui(self):
        # 创建主框架
        main_frame = ttk.Frame(self.window, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 选择文件按钮
        self.select_button = ttk.Button(
            main_frame,
            text="选择Word文档",
            command=self.select_file,
            width=20
        )
        self.select_button.pack(pady=10)
        
        # 文件路径显示
        self.file_label = ttk.Label(main_frame, text="未选择文件", wraplength=500)
        self.file_label.pack(pady=5)
        
        # 语言选择框架
        lang_frame = ttk.LabelFrame(main_frame, text="翻译设置", padding="5")
        lang_frame.pack(fill=tk.X, pady=10, padx=5)
        
        # 目标语言选择
        ttk.Label(lang_frame, text="目标语言：").pack(side=tk.LEFT, padx=5)
        self.target_language = tk.StringVar(value="简体中文")
        self.language_combo = ttk.Combobox(
            lang_frame,
            textvariable=self.target_language,
            values=list(self.translator.supported_languages.keys()),
            state="readonly",
            width=20
        )
        self.language_combo.pack(side=tk.LEFT, padx=5)
        
        # 格式选项
        self.preserve_format = tk.BooleanVar(value=True)
        self.format_check = ttk.Checkbutton(
            lang_frame,
            text="保留原文档格式",
            variable=self.preserve_format
        )
        self.format_check.pack(side=tk.LEFT, padx=20)
        
        # 进度条
        self.progress = Progressbar(
            main_frame,
            orient=tk.HORIZONTAL,
            length=500,
            mode='determinate'
        )
        self.progress.pack(pady=10)
        
        # 状态标签
        self.status_label = ttk.Label(main_frame, text="")
        self.status_label.pack(pady=5)
        
        # 添加控制框架
        control_frame = ttk.LabelFrame(main_frame, text="翻译控制", padding="5")
        control_frame.pack(fill=tk.X, pady=10, padx=5)
        
        # 开始翻译按钮
        self.translate_button = ttk.Button(
            control_frame,
            text="开始翻译",
            command=self.start_translation,
            width=15,
            state=tk.DISABLED
        )
        self.translate_button.pack(side=tk.LEFT, padx=5)
        
        # 暂停/继续按钮
        self.pause_button = ttk.Button(
            control_frame,
            text="暂停",
            command=self.toggle_pause,
            width=15,
            state=tk.DISABLED
        )
        self.pause_button.pack(side=tk.LEFT, padx=5)
        
        # 清理缓存按钮
        self.clean_cache_button = ttk.Button(
            control_frame,
            text="清理缓存",
            command=self.clean_cache_with_confirm,
            width=15
        )
        self.clean_cache_button.pack(side=tk.LEFT, padx=5)
        
        # 状态信息框架
        status_frame = ttk.LabelFrame(main_frame, text="状态信息", padding="5")
        status_frame.pack(fill=tk.X, pady=10, padx=5)
        
        # 进度信息
        self.progress_label = ttk.Label(status_frame, text="进度: 0/0")
        self.progress_label.pack(side=tk.LEFT, padx=5)
        
        # 剩余时间
        self.time_label = ttk.Label(status_frame, text="预计剩余时间: --:--")
        self.time_label.pack(side=tk.LEFT, padx=5)
        
        # 缓存状态
        self.cache_label = ttk.Label(status_frame, text="缓存状态: 无")
        self.cache_label.pack(side=tk.LEFT, padx=5)
        
        # 添加说明文本
        help_text = "支持的语言简体中文、繁體中文、日本語、English、Español、Français、Deutsch、한국어、Русский、Italiano"
        help_label = ttk.Label(main_frame, text=help_text, wraplength=500, foreground="gray")
        help_label.pack(pady=10)

    def toggle_pause(self):
        """切换暂停/继续状态"""
        self.is_paused = not self.is_paused
        if self.is_paused:
            self.pause_button.config(text="继续")
            self.status_label.config(text="翻译已暂停")
        else:
            self.pause_button.config(text="暂停")
            self.status_label.config(text="继续翻译...")

    def update_progress_info(self, current, total):
        """更新进度信息"""
        self.progress_label.config(text=f"进度: {current}/{total}")
        
        if self.translation_start_time and current > 0:
            elapsed_time = time.time() - self.translation_start_time
            avg_time_per_para = elapsed_time / current
            remaining_paras = total - current
            estimated_remaining_time = avg_time_per_para * remaining_paras
            
            # 转换为时分秒格式
            hours = int(estimated_remaining_time // 3600)
            minutes = int((estimated_remaining_time % 3600) // 60)
            seconds = int(estimated_remaining_time % 60)
            
            if hours > 0:
                time_str = f"{hours}小时{minutes}分钟"
            elif minutes > 0:
                time_str = f"{minutes}分钟{seconds}秒"
            else:
                time_str = f"{seconds}秒"
                
            self.time_label.config(text=f"预计剩余时间: {time_str}")

    def update_cache_status(self):
        """更新缓存状态"""
        try:
            cache_dir = os.path.join(os.path.dirname(self.file_path), ".translation_cache")
            if os.path.exists(cache_dir):
                cache_files = len(os.listdir(cache_dir))
                cache_size = sum(os.path.getsize(os.path.join(cache_dir, f)) 
                               for f in os.listdir(cache_dir))
                size_str = f"{cache_size/1024/1024:.1f}MB" if cache_size > 1024*1024 else f"{cache_size/1024:.1f}KB"
                self.cache_label.config(text=f"缓存状态: {cache_files}个文件 ({size_str})")
            else:
                self.cache_label.config(text="缓存状态: 无")
        except:
            self.cache_label.config(text="缓存状态: 未知")

    def clean_cache_with_confirm(self):
        """清理缓存并确认"""
        if messagebox.askyesno("确认", "确定要清理所有缓存文件吗？"):
            self.clean_cache()
            self.update_cache_status()
            messagebox.showinfo("成功", "缓存已清理")

    def get_unique_filename(self, base_path, target_language):
        """获取唯一的文件名"""
        dir_path = os.path.dirname(base_path)
        file_name = os.path.basename(base_path)
        name, ext = os.path.splitext(file_name)
        
        # 基础输出文件名
        base_output = os.path.join(dir_path, f"{name}_translated_{target_language}")
        
        # 如果文件不存在，直接返回基础名称
        if not os.path.exists(f"{base_output}{ext}"):
            return f"{base_output}{ext}"
        
        # 如果文件存在，添加序号
        counter = 1
        while os.path.exists(f"{base_output}_{counter}{ext}"):
            counter += 1
        
        return f"{base_output}_{counter}{ext}"

    def start_translation(self):
        if not hasattr(self, 'file_path'):
            messagebox.showerror("错误", "请先选择文件")
            return
        
        try:
            # 重置计时器和进度
            self.translation_start_time = time.time()
            self.processed_paragraphs = 0
            
            # 启用暂停按钮
            self.pause_button.config(state=tk.NORMAL)
            
            # 获取选择的目标语言
            target_language = self.translator.supported_languages[self.target_language.get()]
            
            # 创建文档处理器
            doc_processor = DocumentProcessor(self.translator)
            
            # 创建缓存文件路径
            cache_dir = os.path.join(os.path.dirname(self.file_path), ".translation_cache")
            if not os.path.exists(cache_dir):
                os.makedirs(cache_dir)
                
            cache_file = os.path.join(
                cache_dir, 
                f"{os.path.basename(self.file_path)}_{target_language}_cache.docx"
            )
            progress_file = os.path.join(
                cache_dir,
                f"{os.path.basename(self.file_path)}_{target_language}_progress.txt"
            )
            
            # 检查是否存在未完成的翻译
            last_index = 0
            doc = Document(self.file_path)
            new_doc = None
            
            # 计算总元素数
            total_elements = doc_processor.count_translatable_elements(doc)
            doc_processor.total_elements = total_elements
            self.progress['maximum'] = total_elements
            
            if os.path.exists(cache_file) and os.path.exists(progress_file):
                with open(progress_file, 'r') as f:
                    last_index = int(f.read().strip() or '0')
                    
                if last_index > 0:
                    response = messagebox.askyesno(
                        "发现未完成翻译",
                        f"发现��次翻译到第 {last_index} 个元素，是否继续上次的翻译？"
                    )
                    if response:
                        new_doc = Document(cache_file)
                    else:
                        last_index = 0
            
            if new_doc is None:
                new_doc = Document()
                
            # 翻译文档内容
            try:
                # 逐个处理文档元素
                for element in doc.element.body:
                    while self.is_paused:
                        self.window.update()
                        time.sleep(0.1)
                        continue
                    
                    if element.tag.endswith('p'):
                        # 处理段落
                        text = element.text.strip()
                        if text:
                            new_para = new_doc.add_paragraph()
                            if self.preserve_format.get():
                                try:
                                    # 复制原始段落的样式
                                    new_para._element = element
                                except:
                                    pass
                            
                            translated_text = self.translator.translate_text(text, target_language)
                            if translated_text:
                                new_para.text = translated_text
                            else:
                                new_para.text = text
                            
                            doc_processor.processed_elements += 1
                        else:
                            new_doc.add_paragraph()
                    
                    elif element.tag.endswith('tbl'):
                        # 处理表格
                        try:
                            # 创建临时文档并添加表格
                            temp_doc = Document()
                            temp_doc._body._element.append(element)
                            source_table = temp_doc.tables[0]
                            doc_processor.translate_table(
                                source_table,
                                new_doc,
                                target_language,
                                self.preserve_format.get()
                            )
                        except Exception as table_error:
                            print(f"处理表格时出错: {str(table_error)}")
                            new_doc.add_paragraph("【表格处理失败】")
                            continue
                    
                    self.update_progress(doc_processor.processed_elements, total_elements)
                
                # 翻译文本框
                text_frame_count = 0
                for shape in doc.inline_shapes:
                    if doc_processor.translate_text_frame(
                        shape,
                        new_doc,
                        target_language
                    ):
                        text_frame_count += 1
                        self.update_progress(doc_processor.processed_elements, total_elements)
                
                if text_frame_count > 0:
                    self.status_label.config(text=f"已翻译 {text_frame_count} 个文本框")
                
                # 翻译眉页脚
                for section in doc.sections:
                    doc_processor.translate_section(
                        section,
                        new_doc.sections[0],  # 假设新文档只有一个section
                        target_language
                    )
                    self.update_progress(doc_processor.processed_elements, total_elements)
                
                # 保存���档（使用新的文件名生成方法）
                output_path = self.get_unique_filename(self.file_path, target_language)
                new_doc.save(output_path)
                
                # 清理缓存文件
                self.clean_cache()
                
                self.status_label.config(text="翻译完成！")
                messagebox.showinfo("成功", f"翻译已完成！\n保存至: {output_path}")
                
            except Exception as e:
                # 保存当前进度
                new_doc.save(cache_file)
                with open(progress_file, 'w') as f:
                    f.write(str(doc_processor.processed_elements))
                raise e
                
        except Exception as e:
            messagebox.showerror("错误", f"翻译过程中出错：{str(e)}")
            self.status_label.config(text="翻译失败")
        
        finally:
            self.pause_button.config(state=tk.DISABLED)
            self.progress['value'] = 0
            self.update_cache_status()

    def select_file(self):
        """选择要翻译的Word文档"""
        file_path = filedialog.askopenfilename(
            title="选择Word文档",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if file_path:
            self.file_path = file_path
            self.file_label.config(text=f"已选择: {file_path}")
            self.translate_button.config(state=tk.NORMAL)

    def diagnose_document(self, doc_path):
        """断文档问题"""
        try:
            # 检查文件大小
            file_size = os.path.getsize(doc_path)
            print(f"文件大小: {file_size} 字节")

            # 尝试以二进制方式读取
            with open(doc_path, 'rb') as f:
                content = f.read()
                # 检查文件头部是否为有效的 Office Open XML 格式
                if not content.startswith(b'PK'):
                    return "文件不是有效的 Office Open XML 格式"

            # 尝试打开文档并分析内容
            doc = Document(doc_path)
            
            # 收集文档信息
            info = {
                "段落数": len(doc.paragraphs),
                "样式数": len(doc.styles),
                "节数": len(doc.sections),
                "表格数": len(doc.tables)
            }
            
            # 检查段落内容
            problematic_paras = []
            for i, para in enumerate(doc.paragraphs):
                try:
                    # 尝试访问段落属性
                    style = para.style
                    format = para.paragraph_format
                    runs = para.runs
                except Exception as e:
                    problematic_paras.append(f"段落 {i+1}: {str(e)}")

            return {
                "基本信息": info,
                "问题段落": problematic_paras
            }

        except Exception as e:
            return f"诊断时发生错误: {str(e)}"

    def run(self):
        """运行应用程序"""
        self.window.mainloop()

    def clean_cache(self):
        """清理所有缓存文件"""
        try:
            cache_dir = os.path.join(os.path.dirname(self.file_path), ".translation_cache")
            if os.path.exists(cache_dir):
                for file in os.listdir(cache_dir):
                    try:
                        os.remove(os.path.join(cache_dir, file))
                    except:
                        pass
                os.rmdir(cache_dir)
        except:
            pass

    def update_progress(self, current, total):
        """更新进度信息"""
        self.progress['value'] = current
        self.progress_label.config(text=f"进度: {current}/{total}")
        self.status_label.config(text=f"正在翻译第 {current}/{total} 个元素...")
        self.window.update()
        
        if self.translation_start_time and current > 0:
            elapsed_time = time.time() - self.translation_start_time
            avg_time_per_element = elapsed_time / current
            remaining_elements = total - current
            estimated_remaining_time = avg_time_per_element * remaining_elements
            
            # 转换为时分秒格式
            hours = int(estimated_remaining_time // 3600)
            minutes = int((estimated_remaining_time % 3600) // 60)
            seconds = int(estimated_remaining_time % 60)
            
            if hours > 0:
                time_str = f"{hours}小时{minutes}分钟"
            elif minutes > 0:
                time_str = f"{minutes}分钟{seconds}秒"
            else:
                time_str = f"{seconds}秒"
                
            self.time_label.config(text=f"预计剩余时间: {time_str}")

def main():
    app = TranslatorGUI()
    app.run()

if __name__ == "__main__":
    main() 